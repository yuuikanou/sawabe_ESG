# coding=UTF-8
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import openpyxl
import numpy as np
import os

import matplotlib
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, RegularPolygon
from matplotlib.path import Path
from matplotlib.projections.polar import PolarAxes
from matplotlib.projections import register_projection
from matplotlib.spines import Spine
from matplotlib.transforms import Affine2D
from matplotlib.lines import Line2D
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Inches, Pt


import shutil


root = tk.Tk()

def open_file():
    filepath = filedialog.askopenfilename()
    entry_path.delete(0, tk.END)
    entry_path.insert(0, filepath)


def open_folder():
    file_path2 = filedialog.askdirectory()
    entry_path2.delete(0, tk.END)
    entry_path2.insert(0, file_path2)

#===============================================================================
#メッセージ出力用のスクリプト
#順位、企業数を引数に入れて、該当するメッセージを出力する

#総合スコア
def m_total(index,l):
    ma = ["貴社のESG対応は中小企業として優れた水準にあります","貴社のESG対応は中小企業として平均的な水準にあります","貴社のESG対応は中小企業としても低い水準にとどまっています"]
    if(index<=l/3):
        m = ma[0]
    elif(l/3<index<=l/3*2):
        m = ma[1]
    else:
        m = ma[2]
    return m
    
#5項目
#同一のため、termを引数に追加
def ma(term,index,l):
    ma = ["ESGに関する"+term+"という側面において中小企業として優れています","ESGに関する"+term+"という側面において中小企業として平均的です","ESGに関する"+term+"という側面において中小企業としても改善が必要な状態です"]
    if(index<=l/3):
        m = ma[0]
    elif(l/3<index<=l/3*2):
        m = ma[1]
    else:
        m = ma[2]
    return m

#14項目
def m1(index,l):
    mst_y=["経営理念にSDGsがしっかり組み込まれています","SDGsの経営理念への反映は中小企業としては平均レベルです","経営理念を策定しSDGsをしっかり組み込んでいくことが必要です"]
    if(index<=l/3):
        m = mst_y[0]
    elif(l/3<index<=l/3*2):
        m = mst_y[1]
    else:
        m = mst_y[2]
    return m
def m2(index,l):
    mst_o=["目標を設定し計画的にとりくむことができています","目標設定や計画的取組は中小企業の平均的レベルです","具体的な目標を設定し、計画的に目標達成に取り組んでいくことが必要です"]
    if(index<=l/3):
        m = mst_o[0]
    elif(l/3<index<=l/3*2):
        m = mst_o[1]
    else:
        m = mst_o[2]
    return m
def m3(index,l):
    mst_g=["SDGsについて組織的な体制がしっかりできています","SDGsに関する組織的体制は中小企業の平均的なレベルです","SDGsについて組織的に取り組む体制の構築が必要です"]
    if(index<=l/3):
        m = mst_g[0]
    elif(l/3<index<=l/3*2):
        m = mst_g[1]
    else:
        m = mst_g[2]
    return m
def m4(index,l):
    mst_w=["SDGsに関する認証取得の水準は高いと評価できます","SDGsに関する認証取得の水準は中小企業として平均的です","SDGsに関する認証取得の水準は中小企業としても低い状態です"]
    if(index<=l/3):
        m = mst_w[0]
    elif(l/3<index<=l/3*2):
        m = mst_w[1]
    else:
        m = mst_w[2]
    return m
def m5(index,l):
    men_y=["経営理念に環境への配慮がしっかり組み込まれています","環境の側面からみた貴社の経営理念は、中小企業として平均的です","環境の側面からみた貴社の経営理念は、中小企業としても平均よりも低い状態です"]
    if(index<=l/3):
        m = men_y[0]
    elif(l/3<index<=l/3*2):
        m = men_y[1]
    else:
        m = men_y[2]
    return m
def m6(index,l):
    men_a=["中小企業として環境改善に対して優れた取組を行っています","中小企業として環境改善への取組は平均的です","中小企業としても環境改善への取組は低い水準です"]
    if(index<=l/3):
        m = men_a[0]
    elif(l/3<index<=l/3*2):
        m = men_a[1]
    else:
        m = men_a[2]
    return m
def m7(index,l):
    men_g=["環境改善にむけた組織的な体制は中小企業として優れた水準です","環境改善にむけた組織的な体制は中小企業として平均的な水準です","環境改善にむけた組織的な体制は中小企業としても低い水準です"]
    if(index<=l/3):
        m = men_g[0]
    elif(l/3<index<=l/3*2):
        m = men_g[1]
    else:
        m = men_g[2]
    return m
def m8(index,l):
    mso_y=["経営理念に社会への配慮がしっかり組み込まれています","社会の側面からみた貴社の経営理念は、中小企業として平均的です","社会の側面からみた貴社の経営理念は、中小企業としても平均よりも低い状態です"]
    if(index<=l/3):
        m = mso_y[0]
    elif(l/3<index<=l/3*2):
        m = mso_y[1]
    else:
        m = mso_y[2]
    return m
def m9(index,l):
    mso_g=["社会という観点からみた組織的な体制は中小企業として優れた水準です","社会という観点からみた組織的な体制は中小企業として平均的な水準です","社会という観点からみた組織的な体制は中小企業としても低い水準です"]
    if(index<=l/3):
        m = mso_g[0]
    elif(l/3<index<=l/3*2):
        m = mso_g[1]
    else:
        m = mso_g[2]
    return m
def m10(index,l):
    mso_a=["社会という観点からみて貴社の取組は中小企業として高い水準にあります","社会という観点からみて貴社の取組は中小企業として平均的です","社会という観点からみて貴社の取組は中小企業としても低い状態です"]
    if(index<=l/3):
        m = mso_a[0]
    elif(l/3<index<=l/3*2):
        m = mso_a[1]
    else:
        m = mso_a[2]
    return m
def m11(index,l):
    mgo_y=["経営理念にガバナンスがしっかり組み込まれています","ガバナンスの側面からみた貴社の経営理念は、中小企業として平均的です","ガバナンスの側面からみた貴社の経営理念は、中小企業としても平均よりも低い状態です"]
    if(index<=l/3):
        m = mgo_y[0]
    elif(l/3<index<=l/3*2):
        m = mgo_y[1]
    else:
        m = mgo_y[2]
    return m
def m12(index,l):
    mgo_g=["ガバナンスという観点からみた組織的な体制は中小企業として優れた水準です","ガバナンスという観点からみた組織的な体制は中小企業として平均的な水準です","ガバナンスという観点からみた組織的な体制は中小企業としても低い水準です"]
    if(index<=l/3):
        m = mgo_g[0]
    elif(l/3<index<=l/3*2):
        m = mgo_g[1]
    else:
        m = mgo_g[2]
    return m
def m13(index,l):
    mar_y=["経営理念に地域の観点がしっかり組み込まれています","地域の観点からみた貴社の経営理念は、中小企業として平均的です","地域の観点からみた貴社の経営理念は、中小企業としても平均よりも低い状態です"]
    if(index<=l/3):
        m = mar_y[0]
    elif(l/3<index<=l/3*2):
        m = mar_y[1]
    else:
        m = mar_y[2]
    return m
def m14(index,l):
    mar_g=["地域社会への貢献という観点からみて組織的な体制がしっかり取り組むことができています","地域社会への貢献という観点からみて組織的体制や組織的取組は中小企業の平均的なレベルです","地域社会への貢献という観点からみて組織的体制を構築し、しっかり取組むことが求められます"]
    if(index<=l/3):
        m = mar_g[0]
    elif(l/3<index<=l/3*2):
        m = mar_g[1]
    else:
        m = mar_g[2]
    return m
    
mlist = [m1,m2,m3,m4,m5,m6,m7,m8,m9,m10,m11,m12,m13,m14]
#=============================================================================-



#==============================================================================
#点数計算を関数として記述
def p4(y):
    if(y == "経営理念がある"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

        st_y += 5
    
    else:
        st_y +=1  

def p5(y):

    
    if(y=="経営理念に反映済み"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_y += 5
        en_y+= 5
        so_y+= 5
        go_y+= 5
        ar_y+= 5
        #print([5,5,5,5,5])
    elif(y=="もともとの経営理念とは別にSDGs等について別の理念がある"):
        
        st_y += 5
        en_y+= 5
        so_y+= 5
        go_y+= 5
        ar_y+= 5
        #print([5,5,5,5,5])
    elif(y=="検討中"):
        
        st_y += 2
        en_y+= 2
        so_y+= 2
        go_y+= 2
        ar_y+= 2
        #print([2,2,2,2,2])
    else:
        
        st_y += 1
        en_y+= 1
        so_y+= 1
        go_y+= 1
        ar_y+= 1
        #print([1,1,1,1,1])

def p6(y):

    if(y=="目標設定している"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_o += 5
        #print([5,0,0,0,0])
    elif(y=="検討中"):
        
        #print([2,0,0,0,0])
        st_o += 2
    elif(y=="目標設定していない"):
        
        st_o += 1
        #print([1,0,0,0,0])
    else:
        pass

def p7(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    sdgs2 = ["貧困をなくそう","飢餓をゼロに","すべての人に健康と福祉を","質の高い教育をみんなに","ジェンダー平等を実現しよう","働きがいも経済成長も","産業と技術革新の基盤を作ろう","人や国の不平等をなくそう","つくる責任、つかう責任","平和と公正をすべての人に","パートナーシップで目標を達成しよう"]
    sdgs3 = ["安全な水とトイレを世界中に","エネルギーをみんなに。そしてクリーンに"]
    sdgs4 = ["気候変動に具体的な対策を","海の豊かさを守ろう","陸の豊かさも守ろう"]
    sdgs5 = ["住み続けられるまちづくりを"]

    for t2 in sdgs2:
        if(t2 in y):
            so_y +=5
            go_y +=5
            
    for t3 in sdgs3:
        if(t3 in y):
            en_y+=5
            so_y +=5
            go_y +=5
            
    for t4 in sdgs4:
        if(t4 in y):
            en_y+=5
            go_y +=5
            
    for t5 in sdgs5:
        if(t4 in y):
            so_y+=5
            go_y+=5
            ar_y+=5
            
def p8(y):
    

    if(y=="経営計画を策定している"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_o += 5
        #print([5,0,0,0,0])
    elif(y=="検討中"):
        
        st_o += 2
    # print([2,0,0,0,0])
    elif(y=="経営計画を策定していない"):
        
        st_o += 1
        #print([1,0,0,0,0])
    else:
        pass

def p8_1(y):
    

    p_list = ["年次","半期","月次","なし"]
    if(p_list[2] in y):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_o += 3
    elif(p_list[1] in y):
        
        st_o += 2
    elif(p_list[0] in y):
        
        st_o += 1
    if(p_list[3] in y):
        
        st_o += 0

def p9(y):
    

    if(y=="経営計画を定期的にモニタリングしている"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_o += 5
    # print([5,0,0,0,0])
    elif(y=="経営計画を定期的にモニタリングしていない"):
        
        st_o += 1
        #print([1,0,0,0,0])
    else:
        pass

def p10(y):
    


    h_list = ["人事･労務担当役員","専門部署","兼任する部署","人事･労務担当者","いない"]
    if(h_list[0] in y):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_g +=1
        so_g +=1
    # print([1,0,1,0,0])
    if(h_list[1] in y):
        
        st_g +=4
        so_g +=4
    #  print([4,0,4,0,0])
    if(h_list[2] in y):
        
        st_g +=3
        so_g +=3
        #print([3,0,3,0,0])
    if(h_list[3] in y):
        
        st_g +=2
        so_g +=2
        #print([2,0,2,0,0])
    if(h_list[4] in y):
        st_g +=1
        so_g +=1
    #print([1,0,1,0,0])
    
def p11(y):


    if(y=="働き方改革を行っている"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        so_g +=5
    # print([0,0,5,0,0])
    elif(y=="検討中"):
        so_g +=2
    # print([0,0,2,0,0])
    elif(y=="働き方改革を行っていない"):
        so_g +=1
    # print([0,0,1,0,0])
    else:
        pass

def p12(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    n_list = ["くるみん","健康経営法人","ユースえーる"]
    for n in n_list:
        if(n in y):
        # print([5,0,5,0,0])
            st_w+=5
            so_a+=5

def p13(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    m_list = ["担当役員","専門部署","兼任する部署","担当者","検討中","いない"]
    if(m_list[0] in y):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_g +=1
        so_g +=1
        go_g +=1
        #print([1,0,1,1,0])
    if(m_list[1] in y):
        st_g +=4
        so_g +=4
        go_g +=4
        #print([4,0,4,4,0])
    if(m_list[2] in y):
        st_g +=3
        so_g +=3
        go_g +=3
        #print([3,0,3,3,0])
    if(m_list[3] in y):
        st_g +=2
        so_g +=2
        go_g +=2
        #print([2,0,2,2,0])
    if(m_list[4] in y):
        st_g +=1
        so_g +=1
        go_g +=1
    # print([1,0,1,1,0])
    if(m_list[5] in y):
        st_g +=1
        so_g +=1
        go_g +=1
    # print([1,0,1,1,0])

def p14(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    net_list = ["ホームページ管理","SNS管理","消費者からの問合せ･クレーム対応","地域イベントへの対応","業界団体での活動","PR冊子の作成"]
    if(y == "回答なし"):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        so_g += 1
        go_g +=1
        #print([0,0,1,1,0])
    else:
        
        for net in net_list:
            if(net in y):
                so_g += 5
                go_g += 5
                #print([0,0,5,5,0])

def p15(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    #注意
    if(y is not None):
        
        if(type(y)==int or type(y)==float):
            if(y==0):
                global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
                so_g+=1
                #print(y,"１点")
            else:
                so_g+=5
                #print(y,"5点")
        else:
            try:
                y_s = int(y[0])
            except:
                y_s = y
            if(y_s != 0 and type(y_s) is int):
                #print(y,"5点")
                so_g+=5
            elif(y_s == 0 and type(y_s) is int):
            #print(y,"1点")
                so_g+=1
            elif(type(y_s) is str):
                so_g+=1
            # print(y,"1点")
    else:
        so_g+=1
        

    #print("まだ")
    
def p16(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    e_list = ["節水","省エネ","クールビズ","ウォームビズ","社外清掃活動","ゴミ削減","土壌汚染対策","排水･水質汚染対策","グリーン調達","環境負荷の低い製品・サービスの開発","有機肥料を使用"]
    if(y == "回答回答なし"):
        en_a += 1
        
    else:
        for e in e_list:
            if(e in y):
                en_a += 5
                
def p17(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    #print("まだ")
    m_list = ["担当役員","専門部署","兼任する部署","担当者","検討中","いない"]
    if(m_list[0] in y):
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g
        st_g +=1
        en_g +=1
        #print([1,1,0,0,0])
    if(m_list[1] in y):
        st_g +=4
        en_g +=4
    # print([4,4,0,0,0])
    if(m_list[2] in y):
        st_g +=3
        en_g +=3
    # print([3,3,0,0,0])
    if(m_list[3] in y):
        st_g +=2
        en_g +=2
    #  print([2,2,0,0,0])
    if(m_list[4] in y):
        st_g +=1
        en_g +=1
        #print([1,1,0,0,0])
    if(m_list[5] in y):
        st_g +=1
        en_g +=1
    # print([1,1,0,0,0])

def p17_1(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    e_list = ["環境負荷の測定","環境関連の情報収集","環境監査","環境関連の顧客やサプライヤからの問合せ対応","環境配慮型製品開発の補助","環境報告書の発行","環境負荷低減のための組織的目標の設定","環境関連の計画設定"]
    if(y is None):
        pass
    else:
        for e in e_list:
            if(e in y):
                en_g+=3
                #print([0,3,0,0,0])
    #print("まだ")

def p19(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    if(y=="ある"):
        st_o += 5
        en_g+=5
    # print([5,5,0,0,0])
    elif(y=="ない"):
        st_o += 1
        en_g+=1
        #print([1,1,0,0,0])
    else:
        pass

def p19_1(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    #注意
    #print("まだ")
    if(y is not None):
        if("無し" in y or"無い" in y or"なし" in y or"ない" in y):
            st_w += 1
            en_g += 1
        else:
            print(y)
            st_w += 5
            en_g += 5
    else:
        st_w += 1
        en_g += 1

def p20(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    if(y=="ある"):
        st_o += 5
        en_g+=5
        #print([5,5,0,0,0])
    elif(y=="検討中"):
        st_o += 2
        en_g+=2
        #print([2,2,0,0,0])
    elif(y=="ない"):
        st_o += 1
        en_g+=1
        #print([1,1,0,0,0])
    else:
        pass

def p22(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    a_list = ["自治体の委員会等への参加","自治体への寄付","地域社会のイベント（祭り等の年中行事や小学校などの行事など）の主催･共催","地域社会の名産品･特産品の製造","ふるさと納税に製品･サービスを提供している","地域社会からの優先雇用","地域の復興","地域の学校との連携","子供の教育支援（奨学金や学習支援など）","子ども食堂の運営や参加","地域の清掃活動","地域企業との連携","離島との取引","地域通貨への対応","地域の環境整備（例えば近隣の植栽の整備）"]

    if("回答なし" in y):
        so_a+=1
        ar_g+=1
        #print([0,0,1,0,1])

    for a in a_list:
        if(a in y):
            so_a+=5
            ar_g+=5
            #print([0,0,5,0,5])

def p23(y):
    global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

    m_list = ["担当役員","専門部署","兼任する部署","担当者","検討中","いない"]
    if(m_list[0] in y):
        st_g +=1
        en_g +=1
        ar_g+=1
        #print([1,0,0,0,1])
    if(m_list[1] in y):
        st_g+=4
        en_g +=4
        ar_g+=4
        #print([4,0,0,0,4])
    if(m_list[2] in y):
        st_g +=3
        en_g +=3
        ar_g+=3
        #print([3,0,0,0,3])
    if(m_list[3] in y):
        st_g +=2
        en_g +=2
        ar_g+=2
        #print([2,0,0,0,2])
    if(m_list[4] in y):
        st_g +=1
        en_g +=1
        ar_g+=1
    #print([1,0,0,0,1])
    if(m_list[5] in y):
        st_g +=1
        en_g +=1
        ar_g +=1
        #print([1,0,0,0,1])

#どこの回答を参考にするかのリスト(実行する関数の引数)
a_list = [4,6,8,9,27,28,29,30,40,39,41,51,59,62,75,84,85,88,87,89,106]

#実行する関数のリスト
p_list = [p4,p5,p6,p7,p8,p8_1,p9,p10,p11,p12,p13,p14,p15,p16,p17,p17_1,p19,p19_1,p20,p22,p23]

#============================================================================

#=======================================================================
#レーダーチャートの雛形
def radar_factory(num_vars, frame='circle'):
    """Create a radar chart with `num_vars` axes.

    This function creates a RadarAxes projection and registers it.

    Parameters
    ----------
    num_vars : int
        Number of variables for radar chart.
    frame : {'circle' | 'polygon'}
        Shape of frame surrounding axes.

    """
    # calculate evenly-spaced axis angles
    theta = np.linspace(0, 2*np.pi, num_vars, endpoint=False)

    class RadarAxes(PolarAxes):

        name = 'radar'

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            # rotate plot such that the first axis is at the top
            self.set_theta_zero_location('N')

        def fill(self, *args, closed=True, **kwargs):
            """Override fill so that line is closed by default"""
            return super().fill(closed=closed, *args, **kwargs)

        def plot(self, *args, **kwargs):
            """Override plot so that line is closed by default"""
            lines = super().plot(*args, **kwargs)
            for line in lines:
                self._close_line(line)

        def _close_line(self, line):
            x, y = line.get_data()
            # FIXME: markers at x[0], y[0] get doubled-up
            if x[0] != x[-1]:
                x = np.concatenate((x, [x[0]]))
                y = np.concatenate((y, [y[0]]))
                line.set_data(x, y)

        def set_varlabels(self, labels, font):
            self.set_thetagrids(np.degrees(theta), labels, fontproperties=font)

        def _gen_axes_patch(self):
            # The Axes patch must be centered at (0.5, 0.5) and of radius 0.5
            # in axes coordinates.
            if frame == 'circle':
                return Circle((0.5, 0.5), 0.5)
            elif frame == 'polygon':
                return RegularPolygon((0.5, 0.5), num_vars,
                                    radius=.5, edgecolor="k")
            else:
                raise ValueError("unknown value for 'frame': %s" % frame)

        def draw(self, renderer):
            """ Draw. If frame is polygon, make gridlines polygon-shaped """
            if frame == 'polygon':
                gridlines = self.yaxis.get_gridlines()
                for gl in gridlines:
                    gl.get_path()._interpolation_steps = num_vars
            super().draw(renderer)


        def _gen_axes_spines(self):
            if frame == 'circle':
                return super()._gen_axes_spines()
            elif frame == 'polygon':
                # spine_type must be 'left'/'right'/'top'/'bottom'/'circle'.
                spine = Spine(axes=self,
                            spine_type='circle',
                            path=Path.unit_regular_polygon(num_vars))
                # unit_regular_polygon gives a polygon of radius 1 centered at
                # (0, 0) but we want a polygon of radius 0.5 centered at (0.5,
                # 0.5) in axes coordinates.
                spine.set_transform(Affine2D().scale(.5).translate(.5, .5)
                                    + self.transAxes)

                return {'polar': spine}
            else:
                raise ValueError("unknown value for 'frame': %s" % frame)

    register_projection(RadarAxes)
    return theta
#=======================================================================

#=======================================================================
#実行ボタンをクリックして実行されるスクリプト
def main():
    filepath = entry_path.get()
    file_path2 = entry_path2.get()
    #filepath = "/Users/kanouyuui/Desktop/ESG評価表サンプル集計結果　レーダーチャート換算方法完成版.xlsx"


#===================================
#エクセルファイルから情報を取得

    # Excelファイルを開く
    workbook = openpyxl.load_workbook(filepath)

    # シートを取得する
    sheetName = dropdown.get()
    sheet = workbook[sheetName]


    rows_with_data = []
    for row in sheet.iter_rows():
        row_data = []
        empty_row = True
        for cell in row:
            row_data.append(cell.value)
            if cell.value is not None:
                empty_row = False
        if not empty_row:
            rows_with_data.append(row_data)
    #print(rows_with_data[0])
#===================================
        
#===================================
#ループを回して、点数計算を行う

    #企業名、各点数が代入されたリスト
    all_list =[]   

    num_coop = 2
    while num_coop < len(rows_with_data):
    #while num_coop < 2:
        coopName = rows_with_data[num_coop][3]
        global st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g

        st_y=0
        st_o=0
        st_g=0
        st_w=0
        en_y=0
        en_a=0
        en_g=0
        so_y=0
        so_g=0
        so_a=0
        go_y=0
        go_g=0
        ar_y=0
        ar_g=0

        #各採点(p4~23)を実行
        num =0
        
        m =""
        while (num < len(a_list)):
            y = rows_with_data[num_coop][a_list[num]]
            #print(p_list[num])
            #print(y)
            m+=(str(y)+"\n")
            p_list[num](y)
            num +=1

        all_list.append([coopName,st_y,st_o,st_g,st_w,en_y,en_a,en_g,so_y,so_g,so_a,go_y,go_g,ar_y,ar_g,st_y+st_o+st_g+st_w,en_y+en_a+en_g,so_y+so_g+so_a,go_y+go_g,ar_y+ar_g,st_y+st_o+st_g+st_w+en_y+en_a+en_g+so_y+so_g+so_a+go_y+go_g+ar_y+ar_g,m])
        st_y+st_o+st_g+st_w+en_y+en_a+en_g+so_y+so_g+so_a+go_y+go_g+ar_y+ar_g
        num_coop+=1

    #print(all_list[0],all_list[1])
#===================================

#===================================
#平均、標準偏差、max、minを計算

    
    arr = np.array(all_list)
    
    #[[企業名],[14項目],[5項目],[total]]
    newarr = np.transpose(arr)
    #print(newarr)
    

    means = []
    i=1
    while i<len(newarr)-1:
        mean = np.mean(newarr[i].astype(int))
        means.append(mean)
        i+=1

    stds = []
    i=1
    while i<len(newarr)-1:
        std = np.std(newarr[i].astype(int))
        stds.append(std)
        i+=1

    maxs = []
    i=1
    while i<len(newarr)-1:
        max = np.max(newarr[i].astype(int))
        maxs.append(max)
        i+=1

    mins = []
    i=1
    while i<len(newarr)-1:
        min = np.min(newarr[i].astype(int))
        mins.append(min)
        i+=1

    i = 0
    while i < len(stds)-1:
        #print(mins[i],means[i],maxs[i])
        i+=1


    m_st = round((means[14]-mins[14])*4/(maxs[14]-mins[14])+1,2)
    m_en = round((means[15]-mins[15])*4/(maxs[15]-mins[15])+1,2)
    m_so = round((means[16]-mins[16])*4/(maxs[16]-mins[16])+1,2)
    m_go = round((means[17]-mins[17])*4/(maxs[17]-mins[17])+1,2)
    m_ar = round((means[18]-mins[18])*4/(maxs[18]-mins[18])+1,2)



#===================================
#各企業に関してループを開始
    in_coop = 0
    while in_coop< len(all_list):
    #while in_coop< 2:
        
        document = Document() 
        document2 = Document() 

        #最初
        c_name = all_list[in_coop][0]
        document.add_heading(c_name+"様", 0)
        document2.add_heading(c_name+"様用資料", 0)
        if (rows_with_data[in_coop+1][0]==None):
            document.add_paragraph("回答日：記載なし")
        else:
            document.add_paragraph("回答日："+rows_with_data[in_coop+1][0].strftime("%Y-%m-%d %H:%M:%S") + '\n')
        
        if (rows_with_data[in_coop+1][0]==None):
            document2.add_paragraph("回答日：記載なし")
        else:
            document2.add_paragraph("回答日："+rows_with_data[in_coop+1][0].strftime("%Y-%m-%d %H:%M:%S") + '\n')

        document.add_heading("ESG評価結果通知表", level=2)
        document.add_paragraph('ESG評価表に回答ありがとうございました。貴社のESG評価は次の通りです。\n')
        
        #総合評価について
        document.add_heading('総合評価：'+str(round((all_list[in_coop][20]-mins[19])*4/(maxs[19]-mins[19])+1,2)), level=1)
        sorted_array = np.sort(newarr[20].astype(int))[::-1]
        

        #print(newarr[20],all_list[i][20])
        total_index = np.where(sorted_array == all_list[in_coop][20])[0][0]
        document.add_paragraph(m_total(total_index,len(all_list)))
        document2.add_heading('総合評価：'+str(round((all_list[in_coop][20]-mins[19])*4/(maxs[19]-mins[19])+1,2))+"   総合順位：" +str(total_index+1)+"/"+str(len(all_list)), level=1)


        #項目別評価について
        document.add_heading('項目別評価')
        st=round((all_list[in_coop][15]-mins[14])*4/(maxs[14]-mins[14])+1,2)
        en=round((all_list[in_coop][16]-mins[15])*4/(maxs[15]-mins[15])+1,2)
        so=round((all_list[in_coop][17]-mins[16])*4/(maxs[16]-mins[16])+1,2)
        go=round((all_list[in_coop][18]-mins[17])*4/(maxs[17]-mins[17])+1,2)
        ar=round((all_list[in_coop][19]-mins[18])*4/(maxs[18]-mins[18])+1,2)
        document.add_paragraph("組織体制："+str(st)+"　"+"環境："+str(en)+"　"+"社会："+str(so)+"　"+"ガバナンス："+str(go)+"　"+"地域："+str(ar)+"　")

        
        #レーダーチャートの作成
        
        data_r = [st,en,so,go,ar]
        data_m = [m_st,m_en,m_so,m_go,m_ar]
        if __name__ == '__main__':
            # 日本語フォントの利用
            if os.name == 'nt':
                font_path = '/System/Library/Fonts/ヒラギノ明朝 ProN.ttc'
            else:
                font_path = '/System/Library/Fonts/ヒラギノ明朝 ProN.ttc'
            font = FontProperties(fname=font_path)

            # ラベルとデータ
            data = [
                ['組織体制', '環境', '社会', 'ガバナンス', '地域'],
                ('', [
                    data_r,
                    data_m,
                ]),
            ]

            # データの色
            colors = ['#FF8095', '#55C500']

            # 外枠の装飾
            matplotlib.rc('axes',edgecolor='white', linewidth=1)

            N = len(data[0])
            theta = radar_factory(N, frame='polygon') # polygon：多角形、circle：円

            spoke_labels = data.pop(0)
            title, case_data = data[0]

            fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(projection='radar'))
            fig.subplots_adjust(top=0.85, bottom=0.05)

            # メモリ線を引く
            ax.set_rgrids([1, 2, 3, 4, 5])

            # 最大最小値の設定
            ax.set_ylim([0, 5])

            # タイトルのセット
            ax.set_title(title,  position=(0.5, 1.1), ha='center', fontproperties=font, fontsize='xx-large')

            # レーダーチャートの色、透明度を設定
            i = 0
    
            for d in case_data:
                #print(d)
                line = ax.plot(theta, d, color=colors[i])       # 枠線の描画
                ax.fill(theta, d, alpha=0.25, color=colors[i])  # 塗りつぶし
                i = i + 1

            # 角のラベルを設定
            ax.set_varlabels(spoke_labels, font)

            # ラベルの表示非表示
            plt.tick_params(labelbottom=True,   # 角ラベル表示
                            labelleft=False)     # メモリラベル非表示

            # 説明の描画

            legend_info = ('貴社', '平均')
            plt.legend(legend_info, loc=(0.7, .95), labelspacing=0.1, fontsize='medium', prop=font)
            

            # 図全体の背景透明度
            fig.patch.set_alpha(0)

            os.makedirs('picture',exist_ok=True)
            fig.savefig("picture/"+c_name+'.png') #ファイルに保存
            document.add_picture('picture/'+c_name+'.png', width=Inches(4))
            plt.close()


        tag = ["組織体制","環境","社会","ガバナンス","地域"]
        tag2 = [["理念","目標","体制","認証"],["理念","取組","体制"],["理念","体制","取組"],["理念","取組"],["理念","取組"]]
        mlist = [[m1,m2,m3,m4],[m5,m6,m7],[m8,m9,m10],[m11,m12],[m13,m14]]
        slist = [[all_list[in_coop][1],all_list[in_coop][2],all_list[in_coop][3],all_list[in_coop][4]],[all_list[in_coop][5],all_list[in_coop][6],all_list[in_coop][7]],[all_list[in_coop][8],all_list[in_coop][9],all_list[in_coop][10]],[all_list[in_coop][11],all_list[in_coop][12]],[all_list[in_coop][13],all_list[in_coop][14]]]
        slist2 = [all_list[in_coop][15],all_list[in_coop][16],all_list[in_coop][17],all_list[in_coop][18],all_list[in_coop][19]]
        arrlist = [[newarr[1],newarr[2],newarr[3],newarr[4]],[newarr[5],newarr[6],newarr[7]],[newarr[8],newarr[9],newarr[10]],[newarr[11],newarr[12]],[newarr[13],newarr[14]]]
        arrlist2 = [newarr[15],newarr[16],newarr[17],newarr[18],newarr[19]]
        i=0
        while  i<len(tag):
            document.add_heading("("+str(i+1)+")"+tag[i], level=2)
            sorted_array = np.sort(arrlist2[i].astype(int))[::-1]
            index = np.where(sorted_array == slist2[i])[0][0]
            document.add_paragraph(ma(tag[i],index,len(all_list)))  

            document2.add_heading("("+str(i+1)+")"+tag[i], level=2)
            document2.add_paragraph("順位："+str(index+1)+"/"+str(len(all_list)))  


            j=0
            while j<len(tag2[i]):
                sorted_array = np.sort(arrlist[i][j].astype(int))[::-1]
                index2 = np.where(sorted_array == slist[i][j])[0][0]
                if len(list(set(sorted_array)))<5:
                    print(tag2[i])
                    document.add_paragraph("　・"+mlist[i][j](index,len(all_list)))
                else:
                    document.add_paragraph("　・"+mlist[i][j](index2,len(all_list)))
                    
                document2.add_paragraph("　・"+tag2[i][j]+" "+"点数："+str(slist[i][j])+"  順位：" + str(index2+1) + "/"+ str(len(all_list)))    



                j+=1
            i+=1
        
        document2.add_paragraph(all_list[in_coop][21])



        os.makedirs(file_path2+"/"+c_name,exist_ok=True)    
        document.save(file_path2+'/'+c_name+'/'+'フィードバックシート'+'.docx')
        document2.save(file_path2+'/'+c_name+'/'+c_name+'.docx')

        in_coop+=1


def get_excel_path():
    excel_path = entry_path.get()
    return excel_path

def get_sheet_names():
    excel_path = get_excel_path()
    wb = openpyxl.load_workbook(excel_path)
    sheet_names = wb.sheetnames
    return sheet_names

def update_dropdown(*args):
    sheet_names = get_sheet_names()
    dropdown['values'] = sheet_names


label = tk.Label(root, text="ファイルパス:")
label.pack()

entry_path = tk.Entry(root)
entry_path.pack()
button_open = tk.Button(root, text="ファイルを開く", command=open_file)
button_open.pack()

# ボタンの作成
button = tk.Button(root, text="更新", command=update_dropdown)
button.pack()

dropdown = ttk.Combobox(root)
dropdown.pack()
# プルダウンメニューの初期設定
dropdown.set("シートを選択してください")


entry_path2 = tk.Entry(root)
entry_path2.pack()
button_open2 = tk.Button(root, text="フォルダーを開く", command=open_folder)
button_open2.pack()

button_do = tk.Button(root, text="実行する", command=main)
button_do.pack()


root.mainloop()